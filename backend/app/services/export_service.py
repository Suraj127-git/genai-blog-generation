from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import logging
from datetime import datetime

logger = logging.getLogger(__name__)


class ExportService:
    @staticmethod
    def export_to_pdf(title: str, content: str, topic: str) -> bytes:
        """Export blog to PDF"""
        try:
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter)
            story = []
            
            styles = getSampleStyleSheet()
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                textColor='#1a1a1a',
                spaceAfter=30,
                alignment=TA_CENTER,
                fontName='Helvetica-Bold'
            )
            
            topic_style = ParagraphStyle(
                'Topic',
                parent=styles['Normal'],
                fontSize=12,
                textColor='#666666',
                spaceAfter=20,
                alignment=TA_CENTER,
                fontName='Helvetica-Oblique'
            )
            
            content_style = ParagraphStyle(
                'Content',
                parent=styles['Normal'],
                fontSize=12,
                textColor='#333333',
                spaceAfter=12,
                alignment=TA_LEFT,
                fontName='Helvetica',
                leading=16
            )
            
            # Add title
            story.append(Paragraph(title, title_style))
            story.append(Spacer(1, 0.2 * inch))
            
            # Add topic
            story.append(Paragraph(f"Topic: {topic}", topic_style))
            story.append(Spacer(1, 0.3 * inch))
            
            # Add content (convert markdown to paragraphs)
            paragraphs = content.split('\n')
            for para in paragraphs:
                if para.strip():
                    # Simple markdown handling
                    if para.startswith('# '):
                        story.append(Paragraph(para[2:], styles['Heading1']))
                    elif para.startswith('## '):
                        story.append(Paragraph(para[3:], styles['Heading2']))
                    elif para.startswith('### '):
                        story.append(Paragraph(para[4:], styles['Heading3']))
                    else:
                        story.append(Paragraph(para, content_style))
                    story.append(Spacer(1, 0.1 * inch))
            
            # Add footer
            footer_style = ParagraphStyle(
                'Footer',
                parent=styles['Normal'],
                fontSize=10,
                textColor='#999999',
                alignment=TA_CENTER,
                fontName='Helvetica'
            )
            story.append(Spacer(1, 0.5 * inch))
            story.append(Paragraph(
                f"Generated on {datetime.now().strftime('%B %d, %Y')}",
                footer_style
            ))
            
            doc.build(story)
            pdf_bytes = buffer.getvalue()
            buffer.close()
            
            logger.info("Exported blog to PDF")
            return pdf_bytes
        except Exception as e:
            logger.error(f"Error exporting to PDF: {e}")
            raise
    
    @staticmethod
    def export_to_docx(title: str, content: str, topic: str) -> bytes:
        """Export blog to DOCX"""
        try:
            doc = Document()
            
            # Add title
            title_para = doc.add_heading(title, level=0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.runs[0]
            title_run.font.size = Pt(24)
            title_run.font.color.rgb = RGBColor(26, 26, 26)
            
            # Add topic
            topic_para = doc.add_paragraph(f"Topic: {topic}")
            topic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            topic_run = topic_para.runs[0]
            topic_run.font.size = Pt(12)
            topic_run.font.italic = True
            topic_run.font.color.rgb = RGBColor(102, 102, 102)
            
            doc.add_paragraph()  # Spacer
            
            # Add content (convert markdown to paragraphs)
            paragraphs = content.split('\n')
            for para in paragraphs:
                if para.strip():
                    # Simple markdown handling
                    if para.startswith('# '):
                        doc.add_heading(para[2:], level=1)
                    elif para.startswith('## '):
                        doc.add_heading(para[3:], level=2)
                    elif para.startswith('### '):
                        doc.add_heading(para[4:], level=3)
                    else:
                        p = doc.add_paragraph(para)
                        p_format = p.paragraph_format
                        p_format.space_after = Pt(12)
            
            # Add footer
            doc.add_paragraph()
            footer_para = doc.add_paragraph(
                f"Generated on {datetime.now().strftime('%B %d, %Y')}"
            )
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_run = footer_para.runs[0]
            footer_run.font.size = Pt(10)
            footer_run.font.color.rgb = RGBColor(153, 153, 153)
            
            # Save to buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            docx_bytes = buffer.getvalue()
            buffer.close()
            
            logger.info("Exported blog to DOCX")
            return docx_bytes
        except Exception as e:
            logger.error(f"Error exporting to DOCX: {e}")
            raise


def get_export_service() -> ExportService:
    """Get export service instance"""
    return ExportService()
